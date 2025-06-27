from pathlib import Path
import pandas as pd
from typing import Union
from docx import Document
from docx.shared import Pt, Inches


def write_doc(df: pd.DataFrame, path: Union[str, Path]) -> None:
    df = df.reset_index()

    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(0)

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.5)
        section.left_margin = section.right_margin = Inches(0.5)

    for row_no, (_, row) in enumerate(df.iterrows(), start=1):
        for col_idx, (col_name, value) in enumerate(row.items(), start=1):
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_after = Pt(0)

            p.add_run(f"{col_name}:").bold = True
            p.add_run("\n")

            text_value = "" if pd.isna(value) else str(value)
            p.add_run(text_value)

            if col_idx < len(row):
                p.add_run("\n")

        if row_no < len(df):
            doc.add_page_break()

    doc.save(str(path))


def iqr_bounds(s: pd.Series, k: float = 1.5) -> tuple[float, float]:
    """
    Return Tukey-fence outlier bounds for a numeric Series.

    Parameters
    ----------
    s : pandas.Series
    k : float, default 1.5
        Multiplier for the IQR (use 3.0 for “extreme” fences).

    Returns
    -------
    (lower_limit, upper_limit)
    """
    q1 = s.quantile(0.25)
    q3 = s.quantile(0.75)
    iqr = q3 - q1
    return q1 - k * iqr, q3 + k * iqr